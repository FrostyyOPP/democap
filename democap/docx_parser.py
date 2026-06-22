"""Phase 1 of parsing: read a .docx into ordered raw text blocks.

This module is deliberately dumb: it only extracts text in document order
(paragraphs + table cells). All interpretation happens in step_extractor.py.
Keeping them separate means we can swap the extractor logic without touching
docx mechanics, and we can unit-test extraction on plain strings.
"""

from __future__ import annotations

import os
from dataclasses import dataclass

from docx import Document


@dataclass
class RawBlock:
    text: str
    style: str          # paragraph style name, e.g. "Heading 1", "List Number"
    is_heading: bool


def _iter_block_items(doc: Document):
    """Yield paragraph text in document order, including text inside tables."""
    for para in doc.paragraphs:
        yield para.style.name if para.style else "Normal", para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para.style.name if para.style else "Normal", para.text


def parse_docx(path: str) -> tuple[str, list[RawBlock]]:
    """Return (inferred_title, ordered list of non-empty RawBlocks)."""
    path = os.path.expanduser(path)
    if not os.path.isfile(path):
        raise FileNotFoundError(f"Script not found: {path}")

    doc = Document(path)
    blocks: list[RawBlock] = []
    title = ""

    for style, text in _iter_block_items(doc):
        text = (text or "").strip()
        if not text:
            continue
        is_heading = style.lower().startswith("title") or style.lower().startswith("heading")
        if not title and is_heading:
            title = text
        blocks.append(RawBlock(text=text, style=style, is_heading=is_heading))

    if not title:
        title = os.path.splitext(os.path.basename(path))[0]

    return title, blocks
